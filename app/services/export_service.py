import os
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from app.models import Book, Chapter
from app.config import settings

async def compile_book(db: AsyncSession, book_id: str, format: str = "docx") -> str:
    """
    Compiles the chapters of the given book into the specified format.
    Returns the file path.
    """
    # Fetch book and chapters
    book = await db.get(Book, book_id)
    if not book:
        raise Exception("Book not found")
        
    result = await db.execute(select(Chapter).filter(Chapter.book_id == book_id).order_by(Chapter.chapter_number))
    chapters = result.scalars().all()
    
    os.makedirs(settings.upload_dir, exist_ok=True)
    
    if format == "txt":
        filename = f"{book.id}_compiled.txt"
        filepath = os.path.join(settings.upload_dir, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"{book.title.upper()}\n\n")
            for ch in chapters:
                f.write(f"Chapter {ch.chapter_number}\n")
                f.write(f"{ch.content}\n\n")
        return filepath
        
    elif format == "docx":
        filename = f"{book.id}_compiled.docx"
        filepath = os.path.join(settings.upload_dir, filename)
        doc = Document()
        doc.add_heading(book.title, 0)
        
        for ch in chapters:
            doc.add_heading(f"Chapter {ch.chapter_number}", level=1)
            if ch.content:
                doc.add_paragraph(ch.content)
            doc.add_page_break()
            
        doc.save(filepath)
        return filepath
        
    else:
        raise Exception(f"Format {format} not supported")
